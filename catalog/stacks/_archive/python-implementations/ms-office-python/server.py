#!/usr/bin/env python3
"""
MS Office MCP Server (Python)
Full read/write support for Microsoft Office documents

GOTCHA - Paragraph Insertion Order:
When inserting multiple paragraphs, use addnext() in FORWARD order, not addprevious():

    # WRONG - results in reversed order:
    for text in reversed(lines):
        para._element.addprevious(new_para._element)

    # CORRECT - maintains order:
    current = anchor_element
    for text in lines:
        new_para = doc.add_paragraph(text)
        current.addnext(new_para._element)
        current = new_para._element  # chain from new element
"""

import json
import re
from pathlib import Path
from typing import Optional
from mcp.server.fastmcp import FastMCP

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import load_workbook, Workbook

mcp = FastMCP("ms-office")


def expand_path(file_path: str) -> Path:
    """Expand ~ and resolve path"""
    return Path(file_path).expanduser().resolve()


def doc_to_markdown(doc: Document) -> str:
    """Convert a Document to markdown"""
    lines = []

    for para in doc.paragraphs:
        style_name = para.style.name.lower() if para.style else ""
        text = para.text.strip()

        if not text:
            lines.append("")
            continue

        # Handle headings
        if "heading 1" in style_name:
            lines.append(f"# {text}")
        elif "heading 2" in style_name:
            lines.append(f"## {text}")
        elif "heading 3" in style_name:
            lines.append(f"### {text}")
        elif "heading 4" in style_name:
            lines.append(f"#### {text}")
        elif "list" in style_name:
            lines.append(f"- {text}")
        else:
            # Check for bold/italic in runs
            formatted = ""
            for run in para.runs:
                t = run.text
                if run.bold and run.italic:
                    formatted += f"***{t}***"
                elif run.bold:
                    formatted += f"**{t}**"
                elif run.italic:
                    formatted += f"*{t}*"
                else:
                    formatted += t
            lines.append(formatted or text)

    # Handle tables
    for table in doc.tables:
        lines.append("")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            lines.append("| " + " | ".join(cells) + " |")
            if i == 0:
                lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
        lines.append("")

    return "\n".join(lines)


@mcp.tool()
def docx_read(path: str, output: Optional[str] = None) -> str:
    """
    Read a Microsoft Word document and extract content as markdown.

    Args:
        path: Path to the .docx file
        output: Optional path to save the markdown output
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        doc = Document(str(file_path))
        markdown = doc_to_markdown(doc)

        result = f"**Document:** {file_path.name}\n\n---\n\n{markdown}"

        if output:
            out_path = expand_path(output)
            out_path.write_text(markdown, encoding="utf-8")
            result = f"Saved to {out_path}\n\n{result}"

        return result
    except Exception as e:
        return f"Error reading document: {e}"


@mcp.tool()
def docx_replace_text(
    path: str,
    find: str,
    replace: str,
    output: Optional[str] = None,
    case_sensitive: bool = True
) -> str:
    """
    Find and replace text in a Word document.

    Args:
        path: Path to the .docx file
        find: Text to find
        replace: Text to replace with
        output: Optional output path (defaults to overwriting original)
        case_sensitive: Whether search is case-sensitive (default True)
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        doc = Document(str(file_path))
        count = 0
        flags = 0 if case_sensitive else re.IGNORECASE

        # Replace in paragraphs
        for para in doc.paragraphs:
            for run in para.runs:
                if re.search(find, run.text, flags):
                    run.text = re.sub(find, replace, run.text, flags=flags)
                    count += 1

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if re.search(find, run.text, flags):
                                run.text = re.sub(find, replace, run.text, flags=flags)
                                count += 1

        # Save
        out_path = expand_path(output) if output else file_path
        doc.save(str(out_path))

        return f"Replaced {count} occurrence(s) of '{find}' with '{replace}'\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def docx_edit_table(
    path: str,
    table_index: int,
    row: int,
    col: int,
    value: str,
    output: Optional[str] = None
) -> str:
    """
    Edit a specific cell in a table.

    Args:
        path: Path to the .docx file
        table_index: Which table (0-indexed)
        row: Row number (0-indexed)
        col: Column number (0-indexed)
        value: New cell value
        output: Optional output path (defaults to overwriting original)
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        doc = Document(str(file_path))

        if table_index >= len(doc.tables):
            return f"Error: Table {table_index} not found. Document has {len(doc.tables)} table(s)."

        table = doc.tables[table_index]

        if row >= len(table.rows):
            return f"Error: Row {row} not found. Table has {len(table.rows)} row(s)."

        if col >= len(table.rows[row].cells):
            return f"Error: Column {col} not found. Row has {len(table.rows[row].cells)} column(s)."

        old_value = table.rows[row].cells[col].text
        table.rows[row].cells[col].text = value

        out_path = expand_path(output) if output else file_path
        doc.save(str(out_path))

        return f"Updated table[{table_index}] cell[{row},{col}]: '{old_value}' -> '{value}'\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def docx_update_table_column(
    path: str,
    table_index: int,
    old_header: str,
    new_header: str,
    output: Optional[str] = None
) -> str:
    """
    Rename a table column header.

    Args:
        path: Path to the .docx file
        table_index: Which table (0-indexed)
        old_header: Current header text to find
        new_header: New header text
        output: Optional output path
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        doc = Document(str(file_path))

        if table_index >= len(doc.tables):
            return f"Error: Table {table_index} not found."

        table = doc.tables[table_index]
        header_row = table.rows[0]

        found = False
        for cell in header_row.cells:
            if cell.text.strip() == old_header:
                cell.text = new_header
                found = True
                break

        if not found:
            headers = [c.text.strip() for c in header_row.cells]
            return f"Error: Header '{old_header}' not found. Available: {headers}"

        out_path = expand_path(output) if output else file_path
        doc.save(str(out_path))

        return f"Renamed column header: '{old_header}' -> '{new_header}'\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def docx_delete_table_column(
    path: str,
    table_index: int,
    col: int,
    output: Optional[str] = None
) -> str:
    """
    Delete a column from a table.

    Args:
        path: Path to the .docx file
        table_index: Which table (0-indexed)
        col: Column index to delete (0-indexed)
        output: Optional output path
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        doc = Document(str(file_path))

        if table_index >= len(doc.tables):
            return f"Error: Table {table_index} not found."

        table = doc.tables[table_index]

        # Get the column header for reporting
        header = table.rows[0].cells[col].text if col < len(table.rows[0].cells) else f"column {col}"

        # Delete cell from each row
        for row in table.rows:
            if col < len(row.cells):
                cell = row.cells[col]
                # Remove the cell's XML element
                cell._tc.getparent().remove(cell._tc)

        out_path = expand_path(output) if output else file_path
        doc.save(str(out_path))

        return f"Deleted column {col} ('{header}')\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def docx_list_tables(path: str) -> str:
    """
    List all tables in a document with their structure.

    Args:
        path: Path to the .docx file
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        doc = Document(str(file_path))

        if not doc.tables:
            return "No tables found in document."

        result = f"**Document:** {file_path.name}\n\n"
        result += f"**Tables:** {len(doc.tables)}\n\n"

        for i, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            result += f"### Table {i}\n"
            result += f"- Rows: {len(table.rows)}\n"
            result += f"- Columns: {len(headers)}\n"
            result += f"- Headers: {headers}\n\n"

            # Show preview
            result += "| " + " | ".join(headers) + " |\n"
            result += "| " + " | ".join(["---"] * len(headers)) + " |\n"
            for row in table.rows[1:4]:  # First 3 data rows
                cells = [cell.text.strip()[:20] for cell in row.cells]
                result += "| " + " | ".join(cells) + " |\n"
            if len(table.rows) > 4:
                result += f"*...and {len(table.rows) - 4} more rows*\n"
            result += "\n"

        return result
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def docx_insert_paragraphs(
    path: str,
    after_text: str,
    paragraphs: list[str],
    output: Optional[str] = None
) -> str:
    """
    Insert paragraphs after a specific text marker in the document.

    Args:
        path: Path to the .docx file
        after_text: Text to find - new paragraphs will be inserted after this paragraph
        paragraphs: List of paragraph texts to insert (in order)
        output: Optional output path (defaults to overwriting original)

    Note: Uses addnext() chaining to maintain correct order.
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        doc = Document(str(file_path))

        # Find the anchor paragraph
        anchor_elem = None
        for para in doc.paragraphs:
            if after_text in para.text:
                anchor_elem = para._element
                break

        if anchor_elem is None:
            return f"Error: Could not find paragraph containing '{after_text}'"

        # Insert paragraphs in order using addnext() chaining
        current = anchor_elem
        for text in paragraphs:
            new_para = doc.add_paragraph(text)
            current.addnext(new_para._element)
            current = new_para._element

        out_path = expand_path(output) if output else file_path
        doc.save(str(out_path))

        return f"Inserted {len(paragraphs)} paragraph(s) after '{after_text[:30]}...'\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def docx_create(
    path: str,
    title: str,
    content: Optional[str] = None
) -> str:
    """
    Create a new Word document.

    Args:
        path: Path for the new .docx file
        title: Document title (added as Heading 1)
        content: Optional markdown-like content to add
    """
    file_path = expand_path(path)

    try:
        doc = Document()

        # Add title
        doc.add_heading(title, level=1)

        if content:
            # Simple markdown parsing
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                elif line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

        doc.save(str(file_path))
        return f"Created document: {file_path}"
    except Exception as e:
        return f"Error: {e}"


# =============================================================================
# EXCEL TOOLS
# =============================================================================

@mcp.tool()
def xlsx_read(
    path: str,
    sheet: Optional[str] = None,
    range: Optional[str] = None,
    output: Optional[str] = None
) -> str:
    """
    Read data from an Excel spreadsheet.

    Args:
        path: Path to the .xlsx file
        sheet: Sheet name (default: active sheet)
        range: Cell range like "A1:D10" (default: all data)
        output: Optional path to save as CSV or JSON
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        wb = load_workbook(str(file_path), read_only=True, data_only=True)

        # Get sheet
        if sheet:
            if sheet not in wb.sheetnames:
                return f"Error: Sheet '{sheet}' not found. Available: {wb.sheetnames}"
            ws = wb[sheet]
        else:
            ws = wb.active

        # Get data
        if range:
            data = [[cell.value for cell in row] for row in ws[range]]
        else:
            data = [[cell.value for cell in row] for row in ws.iter_rows()]

        # Convert to markdown table
        result = f"**Spreadsheet:** {file_path.name}\n"
        result += f"**Sheet:** {ws.title}\n"
        result += f"**Sheets available:** {wb.sheetnames}\n\n"

        if data and data[0]:
            # Header
            headers = [str(h) if h else "" for h in data[0]]
            result += "| " + " | ".join(headers) + " |\n"
            result += "| " + " | ".join(["---"] * len(headers)) + " |\n"

            # Data rows
            for row in data[1:]:
                cells = [str(c) if c is not None else "" for c in row]
                # Pad if needed
                while len(cells) < len(headers):
                    cells.append("")
                result += "| " + " | ".join(cells[:len(headers)]) + " |\n"

        wb.close()

        if output:
            out_path = expand_path(output)
            out_path.write_text(result, encoding="utf-8")
            return f"Saved to {out_path}\n\n{result}"

        return result
    except Exception as e:
        return f"Error reading spreadsheet: {e}"


@mcp.tool()
def xlsx_list_sheets(path: str) -> str:
    """
    List all sheets in an Excel workbook.

    Args:
        path: Path to the .xlsx file
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        wb = load_workbook(str(file_path), read_only=True)
        sheets = wb.sheetnames
        wb.close()

        result = f"**Spreadsheet:** {file_path.name}\n\n"
        result += "**Sheets:**\n"
        for i, name in enumerate(sheets):
            result += f"  {i+1}. {name}\n"

        return result
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def xlsx_write_cell(
    path: str,
    sheet: str,
    cell: str,
    value: str,
    output: Optional[str] = None
) -> str:
    """
    Write a value to a specific cell.

    Args:
        path: Path to the .xlsx file
        sheet: Sheet name
        cell: Cell reference like "A1" or "B5"
        value: Value to write
        output: Optional output path (defaults to overwriting original)
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        wb = load_workbook(str(file_path))

        if sheet not in wb.sheetnames:
            return f"Error: Sheet '{sheet}' not found. Available: {wb.sheetnames}"

        ws = wb[sheet]
        old_value = ws[cell].value
        ws[cell] = value

        out_path = expand_path(output) if output else file_path
        wb.save(str(out_path))
        wb.close()

        return f"Updated {sheet}!{cell}: '{old_value}' -> '{value}'\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def xlsx_write_row(
    path: str,
    sheet: str,
    row: int,
    values: list[str],
    start_col: str = "A",
    output: Optional[str] = None
) -> str:
    """
    Write values to a row.

    Args:
        path: Path to the .xlsx file
        sheet: Sheet name
        row: Row number (1-indexed)
        values: List of values to write
        start_col: Starting column letter (default: "A")
        output: Optional output path
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        wb = load_workbook(str(file_path))

        if sheet not in wb.sheetnames:
            return f"Error: Sheet '{sheet}' not found."

        ws = wb[sheet]
        start_col_idx = ord(start_col.upper()) - ord('A') + 1

        for i, val in enumerate(values):
            ws.cell(row=row, column=start_col_idx + i, value=val)

        out_path = expand_path(output) if output else file_path
        wb.save(str(out_path))
        wb.close()

        end_col = chr(ord(start_col.upper()) + len(values) - 1)
        return f"Wrote {len(values)} values to {sheet}!{start_col}{row}:{end_col}{row}\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def xlsx_append_row(
    path: str,
    sheet: str,
    values: list[str],
    output: Optional[str] = None
) -> str:
    """
    Append a row to the end of the data.

    Args:
        path: Path to the .xlsx file
        sheet: Sheet name
        values: List of values for the new row
        output: Optional output path
    """
    file_path = expand_path(path)

    if not file_path.exists():
        return f"Error: File not found: {file_path}"

    try:
        wb = load_workbook(str(file_path))

        if sheet not in wb.sheetnames:
            return f"Error: Sheet '{sheet}' not found."

        ws = wb[sheet]
        ws.append(values)
        new_row = ws.max_row

        out_path = expand_path(output) if output else file_path
        wb.save(str(out_path))
        wb.close()

        return f"Appended row {new_row} with {len(values)} values\nSaved to: {out_path}"
    except Exception as e:
        return f"Error: {e}"


@mcp.tool()
def xlsx_create(
    path: str,
    sheet_name: str = "Sheet1",
    headers: Optional[list[str]] = None
) -> str:
    """
    Create a new Excel workbook.

    Args:
        path: Path for the new .xlsx file
        sheet_name: Name for the first sheet (default: "Sheet1")
        headers: Optional list of column headers
    """
    file_path = expand_path(path)

    try:
        wb = Workbook()
        ws = wb.active
        ws.title = sheet_name

        if headers:
            ws.append(headers)

        wb.save(str(file_path))
        wb.close()

        return f"Created spreadsheet: {file_path}"
    except Exception as e:
        return f"Error: {e}"


if __name__ == "__main__":
    mcp.run()

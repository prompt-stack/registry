"""
Word Document (.docx) operations using python-docx

GOTCHA - Paragraph Insertion Order:
When inserting multiple paragraphs, use addnext() in FORWARD order, not addprevious():

    # WRONG - results in reversed order:
    for text in reversed(lines):
        para._element.addprevious(new_para._element)

    # CORRECT - maintains order:
    current = anchor_element
    for text in lines:
        new_para = doc.add_paragraph(text)
        current.addnext(new_para._element)
        current = new_para._element  # chain from new element
"""

from pathlib import Path
from typing import Optional, List
import re

from docx import Document
from docx.shared import Inches, Pt


class WordDocument:
    """Operations for Word documents"""

    def __init__(self, path: str):
        self.path = Path(path).expanduser().resolve()
        self._doc = None

    def _load(self):
        if not self._doc:
            self._doc = Document(str(self.path))
        return self._doc

    def _save(self, output: Optional[str] = None):
        out_path = Path(output).expanduser().resolve() if output else self.path
        self._doc.save(str(out_path))
        return out_path

    @staticmethod
    def create(path: str, title: str, content: Optional[str] = None) -> 'WordDocument':
        """Create a new Word document"""
        file_path = Path(path).expanduser().resolve()
        doc = Document()

        doc.add_heading(title, level=1)

        if content:
            for line in content.split("\n"):
                line = line.strip()
                if not line:
                    continue
                if line.startswith("### "):
                    doc.add_heading(line[4:], level=3)
                elif line.startswith("## "):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith("# "):
                    doc.add_heading(line[2:], level=1)
                elif line.startswith("- ") or line.startswith("* "):
                    doc.add_paragraph(line[2:], style="List Bullet")
                else:
                    doc.add_paragraph(line)

        doc.save(str(file_path))
        return WordDocument(str(file_path))

    def to_markdown(self) -> str:
        """Convert document to markdown"""
        doc = self._load()
        lines = []

        for para in doc.paragraphs:
            style_name = para.style.name.lower() if para.style else ""
            text = para.text.strip()

            if not text:
                lines.append("")
                continue

            if "heading 1" in style_name:
                lines.append(f"# {text}")
            elif "heading 2" in style_name:
                lines.append(f"## {text}")
            elif "heading 3" in style_name:
                lines.append(f"### {text}")
            elif "heading 4" in style_name:
                lines.append(f"#### {text}")
            elif "list" in style_name:
                lines.append(f"- {text}")
            else:
                formatted = ""
                for run in para.runs:
                    t = run.text
                    if run.bold and run.italic:
                        formatted += f"***{t}***"
                    elif run.bold:
                        formatted += f"**{t}**"
                    elif run.italic:
                        formatted += f"*{t}*"
                    else:
                        formatted += t
                lines.append(formatted or text)

        # Handle tables
        for table in doc.tables:
            lines.append("")
            for i, row in enumerate(table.rows):
                cells = [cell.text.strip() for cell in row.cells]
                lines.append("| " + " | ".join(cells) + " |")
                if i == 0:
                    lines.append("| " + " | ".join(["---"] * len(cells)) + " |")
            lines.append("")

        return "\n".join(lines)

    def replace_text(self, find: str, replace: str, case_sensitive: bool = True) -> int:
        """Find and replace text, returns count of replacements"""
        doc = self._load()
        count = 0
        flags = 0 if case_sensitive else re.IGNORECASE

        for para in doc.paragraphs:
            for run in para.runs:
                if re.search(find, run.text, flags):
                    run.text = re.sub(find, replace, run.text, flags=flags)
                    count += 1

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for run in para.runs:
                            if re.search(find, run.text, flags):
                                run.text = re.sub(find, replace, run.text, flags=flags)
                                count += 1

        return count

    def get_tables(self) -> List[dict]:
        """Get info about all tables"""
        doc = self._load()
        tables = []

        for i, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells]
            tables.append({
                'index': i,
                'rows': len(table.rows),
                'columns': len(headers),
                'headers': headers,
            })

        return tables

    def edit_table_cell(self, table_index: int, row: int, col: int, value: str) -> str:
        """Edit a specific cell in a table"""
        doc = self._load()

        if table_index >= len(doc.tables):
            raise ValueError(f"Table {table_index} not found. Document has {len(doc.tables)} table(s).")

        table = doc.tables[table_index]

        if row >= len(table.rows):
            raise ValueError(f"Row {row} not found. Table has {len(table.rows)} row(s).")

        if col >= len(table.rows[row].cells):
            raise ValueError(f"Column {col} not found.")

        old_value = table.rows[row].cells[col].text
        table.rows[row].cells[col].text = value

        return old_value

    def delete_table_column(self, table_index: int, col: int) -> str:
        """Delete a column from a table"""
        doc = self._load()

        if table_index >= len(doc.tables):
            raise ValueError(f"Table {table_index} not found.")

        table = doc.tables[table_index]
        header = table.rows[0].cells[col].text if col < len(table.rows[0].cells) else f"column {col}"

        for row in table.rows:
            if col < len(row.cells):
                cell = row.cells[col]
                cell._tc.getparent().remove(cell._tc)

        return header

    def insert_paragraphs(self, after_text: str, paragraphs: List[str]) -> int:
        """Insert paragraphs after a text marker, returns count inserted"""
        doc = self._load()

        anchor_elem = None
        for para in doc.paragraphs:
            if after_text in para.text:
                anchor_elem = para._element
                break

        if anchor_elem is None:
            raise ValueError(f"Could not find paragraph containing '{after_text}'")

        current = anchor_elem
        for text in paragraphs:
            new_para = doc.add_paragraph(text)
            current.addnext(new_para._element)
            current = new_para._element

        return len(paragraphs)

    def save(self, output: Optional[str] = None) -> Path:
        """Save the document"""
        return self._save(output)
